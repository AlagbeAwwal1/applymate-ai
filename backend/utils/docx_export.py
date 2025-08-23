from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def markdown_to_docx(md: str, title: str = None) -> Document:
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    if title:
        h = doc.add_heading(title, level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for line in md.splitlines():
        if line.strip().startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(line.strip()[2:])
        else:
            doc.add_paragraph(line)
    return doc
